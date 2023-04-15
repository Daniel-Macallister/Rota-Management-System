def create_docx(title,dates,names,filename):
    from docx import Document
    from docx.shared import Pt
    from docx.shared import RGBColor
    document = Document()
    document.add_picture('logo.png')
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(20)
    head = document.add_paragraph().add_run(title)
    title_font = head.font
    title_font.size = Pt(28)
    title_font.color.rgb = RGBColor(0x29, 0xB3, 0xD9)
    table = document.add_table(rows=16,cols=2)
    table.style = 'Table Grid'
    for i in range(0,16):
        name_cell = table.cell(i,1)
        name_cell.text = str(names[i])
        date_cell = table.cell(i,0)
        date_cell.text = str(dates[i])
    font.size = Pt(14)
    small_text = document.add_paragraph("If you find you are unavailble on your date, please arrange for someone else to cover for you on that day and let the office know.")
    document.save(filename + ".docx")
    
